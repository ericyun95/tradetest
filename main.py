#!/usr/bin/env python3
"""
Global Trade Mentor AI (MVP)
- HS Code 후보 탐색 (로컬 키워드 매칭)
- UN Comtrade+ API로 국가별 수입 데이터 수집 (무료 플랜)
- 수입규모 + 성장률 스코어링 → 타겟 국가 TOP 3 출력

무료 API 키 발급:
  1. https://comtradeplus.un.org 에서 회원가입
  2. My Account → API Keys → 키 복사
  3. 환경변수 설정: export COMTRADE_API_KEY="your_key"
     또는 실행 시 직접 입력
"""

import os
import sys
import time
import datetime
import requests
from difflib import SequenceMatcher
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ──────────────────────────────────────────────
# 로컬 HS Code 데이터베이스
# ──────────────────────────────────────────────
HS_CODE_DB = [
    # 전자/IT
    {"name": "스마트폰",     "code": "851712", "desc": "스마트폰 / Smartphones",
     "kw": ["스마트폰", "핸드폰", "휴대폰", "모바일", "smartphone", "mobile phone"]},
    {"name": "노트북",       "code": "847130", "desc": "노트북 컴퓨터 / Laptop computers",
     "kw": ["노트북", "랩탑", "laptop", "notebook"]},
    {"name": "TV",           "code": "852872", "desc": "LCD/LED TV / Television receivers",
     "kw": ["tv", "텔레비전", "television", "lcd", "led tv"]},
    {"name": "에어컨",       "code": "841510", "desc": "에어컨 / Air conditioners",
     "kw": ["에어컨", "냉방기", "air conditioner", "air conditioning"]},
    {"name": "세탁기",       "code": "845011", "desc": "세탁기 / Washing machines",
     "kw": ["세탁기", "washing machine", "laundry"]},
    {"name": "냉장고",       "code": "841821", "desc": "냉장고 / Refrigerators",
     "kw": ["냉장고", "refrigerator", "fridge"]},
    {"name": "반도체",       "code": "854231", "desc": "반도체 / Electronic integrated circuits",
     "kw": ["반도체", "집적회로", "ic", "semiconductor", "chip", "메모리"]},
    {"name": "배터리",       "code": "850760", "desc": "리튬이온 배터리 / Li-ion batteries",
     "kw": ["배터리", "전지", "리튬이온", "battery", "lithium", "2차전지"]},
    {"name": "카메라",       "code": "900640", "desc": "디지털 카메라 / Digital cameras",
     "kw": ["카메라", "디지털카메라", "camera", "digital camera"]},
    {"name": "이어폰",       "code": "851830", "desc": "이어폰·헤드폰 / Headphones & earphones",
     "kw": ["이어폰", "헤드폰", "earphones", "headphones", "earbuds", "무선이어폰"]},
    {"name": "스피커",       "code": "851840", "desc": "스피커 / Loudspeakers",
     "kw": ["스피커", "speaker", "블루투스스피커", "bluetooth speaker"]},
    {"name": "태블릿",       "code": "847192", "desc": "태블릿 PC / Tablet computers",
     "kw": ["태블릿", "tablet", "ipad", "갤럭시탭"]},
    {"name": "스마트워치",   "code": "851762", "desc": "스마트워치 / Smartwatches",
     "kw": ["스마트워치", "smartwatch", "갤럭시워치", "웨어러블"]},
    {"name": "OLED패널",     "code": "854370", "desc": "OLED 패널 / OLED panels",
     "kw": ["oled", "디스플레이", "패널", "display panel"]},
    {"name": "로봇청소기",   "code": "850910", "desc": "로봇청소기 / Robotic vacuum cleaners",
     "kw": ["로봇청소기", "robot vacuum", "청소로봇"]},
    # 의류/패션
    {"name": "티셔츠",       "code": "610910", "desc": "면 티셔츠 / T-shirts of cotton",
     "kw": ["티셔츠", "t-shirt", "tshirt", "반팔"]},
    {"name": "청바지",       "code": "620342", "desc": "청바지 / Denim trousers",
     "kw": ["청바지", "데님", "jeans", "denim"]},
    {"name": "운동화",       "code": "640411", "desc": "운동화 / Sports footwear",
     "kw": ["운동화", "스포츠화", "sneakers", "sports shoes", "athletic shoes"]},
    {"name": "가방",         "code": "420222", "desc": "핸드백 / Handbags",
     "kw": ["가방", "핸드백", "백", "handbag", "bag", "purse"]},
    {"name": "패딩",         "code": "621300", "desc": "다운재킷 / Down jackets",
     "kw": ["패딩", "다운재킷", "점퍼", "down jacket", "puffer"]},
    {"name": "레깅스",       "code": "611241", "desc": "레깅스 / Leggings",
     "kw": ["레깅스", "leggings", "요가팬츠", "스포츠웨어"]},
    # 식품/음료
    {"name": "커피",         "code": "090111", "desc": "커피 원두 / Coffee beans",
     "kw": ["커피", "원두", "coffee beans", "coffee"]},
    {"name": "녹차",         "code": "090210", "desc": "녹차 / Green tea",
     "kw": ["녹차", "차", "green tea", "tea"]},
    {"name": "홍차",         "code": "090240", "desc": "홍차 / Black tea",
     "kw": ["홍차", "black tea", "발효차"]},
    {"name": "라면",         "code": "190230", "desc": "라면 / Instant noodles",
     "kw": ["라면", "인스턴트면", "noodles", "ramen", "instant noodles"]},
    {"name": "김치",         "code": "200590", "desc": "김치 / Kimchi (prepared vegetables)",
     "kw": ["김치", "kimchi"]},
    {"name": "초콜릿",       "code": "180620", "desc": "초콜릿 / Chocolate",
     "kw": ["초콜릿", "초코", "chocolate", "choco"]},
    {"name": "인삼",         "code": "121120", "desc": "인삼 / Ginseng roots",
     "kw": ["인삼", "홍삼", "ginseng", "red ginseng", "인삼뿌리"]},
    {"name": "김",           "code": "121221", "desc": "김 / Dried seaweed (laver)",
     "kw": ["김", "seaweed", "laver", "nori", "건조김"]},
    {"name": "고추장",       "code": "210390", "desc": "고추장·된장 / Gochujang, sauces",
     "kw": ["고추장", "된장", "간장", "쌈장", "gochujang", "doenjang", "soy sauce", "소스"]},
    {"name": "과자",         "code": "190590", "desc": "과자·비스킷 / Biscuits & cookies",
     "kw": ["과자", "비스킷", "쿠키", "스낵", "biscuit", "cookie", "snack", "crackers"]},
    {"name": "음료",         "code": "220299", "desc": "음료 / Non-alcoholic beverages",
     "kw": ["음료", "음료수", "이온음료", "에너지드링크", "beverage", "drink", "energy drink"]},
    {"name": "소주",         "code": "220860", "desc": "소주 / Korean spirits (Soju)",
     "kw": ["소주", "soju", "증류주", "spirits"]},
    {"name": "막걸리",       "code": "220600", "desc": "막걸리 / Makgeolli (rice wine)",
     "kw": ["막걸리", "makgeolli", "rice wine", "탁주"]},
    {"name": "맥주",         "code": "220300", "desc": "맥주 / Beer",
     "kw": ["맥주", "beer", "lager", "ale"]},
    {"name": "아이스크림",   "code": "210500", "desc": "아이스크림 / Ice cream",
     "kw": ["아이스크림", "ice cream", "gelato", "아이스"]},
    {"name": "유자차",       "code": "200980", "desc": "유자차·과일음료 / Citrus juice beverages",
     "kw": ["유자차", "유자", "yuja", "citrus tea", "fruit tea"]},
    {"name": "떡",           "code": "190190", "desc": "떡·쌀가공식품 / Rice cakes",
     "kw": ["떡", "떡볶이", "rice cake", "tteok", "tteokbokki"]},
    {"name": "냉동식품",     "code": "160520", "desc": "냉동 해산물·가공식품 / Frozen prepared foods",
     "kw": ["냉동식품", "만두", "냉동만두", "frozen food", "dumpling", "gyoza", "mandu"]},
    {"name": "참기름",       "code": "151590", "desc": "참기름·들기름 / Sesame oil",
     "kw": ["참기름", "들기름", "sesame oil", "perilla oil"]},
    # 화장품/뷰티
    {"name": "립스틱",       "code": "330420", "desc": "립스틱 / Lip make-up",
     "kw": ["립스틱", "립글로스", "lipstick", "lip gloss", "립밤", "lip balm"]},
    {"name": "스킨케어",     "code": "330499", "desc": "스킨케어 / Skin care preparations",
     "kw": ["스킨케어", "화장품", "크림", "로션", "스킨", "에센스", "앰플", "세럼",
            "skincare", "cosmetics", "cream", "lotion", "essence", "serum", "ampoule"]},
    {"name": "샴푸",         "code": "330510", "desc": "샴푸 / Shampoos",
     "kw": ["샴푸", "shampoo", "hair wash"]},
    {"name": "향수",         "code": "330300", "desc": "향수 / Perfumes & fragrances",
     "kw": ["향수", "perfume", "fragrance", "cologne"]},
    {"name": "마스크팩",     "code": "330499", "desc": "마스크팩 / Sheet masks",
     "kw": ["마스크팩", "시트마스크", "sheet mask", "face mask pack", "팩"]},
    {"name": "선크림",       "code": "330499", "desc": "선크림·자외선차단제 / Sunscreen",
     "kw": ["선크림", "선스크린", "자외선차단", "sunscreen", "spf", "sunblock"]},
    {"name": "BB크림",       "code": "330420", "desc": "BB크림·파운데이션 / BB cream, foundation",
     "kw": ["bb크림", "cc크림", "파운데이션", "bb cream", "cc cream", "foundation", "쿠션"]},
    {"name": "헤어케어",     "code": "330590", "desc": "헤어케어 / Hair care products",
     "kw": ["헤어에센스", "헤어오일", "트리트먼트", "hair oil", "hair treatment", "conditioner", "컨디셔너"]},
    {"name": "바디로션",     "code": "330499", "desc": "바디로션·바디케어 / Body lotion",
     "kw": ["바디로션", "바디크림", "body lotion", "body cream", "body wash", "바디워시"]},
    # 자동차
    {"name": "자동차",       "code": "870323", "desc": "승용차 / Passenger cars (1500-3000cc)",
     "kw": ["자동차", "승용차", "차", "car", "automobile", "vehicle", "sedan"]},
    {"name": "전기차",       "code": "870380", "desc": "전기차 / Electric vehicles",
     "kw": ["전기차", "ev", "electric vehicle", "electric car", "전기자동차"]},
    {"name": "자동차부품",   "code": "870899", "desc": "자동차 부품 / Auto parts",
     "kw": ["자동차부품", "부품", "auto parts", "car parts"]},
    {"name": "타이어",       "code": "401110", "desc": "타이어 / Pneumatic tyres",
     "kw": ["타이어", "tyre", "tire"]},
    # 에너지/소재
    {"name": "태양광패널",   "code": "854140", "desc": "태양광 패널 / Photovoltaic cells",
     "kw": ["태양광", "솔라", "solar panel", "photovoltaic", "태양전지"]},
    {"name": "철강",         "code": "720839", "desc": "열연 강판 / Hot-rolled steel",
     "kw": ["철강", "철판", "강판", "steel", "iron"]},
    {"name": "플라스틱",     "code": "390760", "desc": "플라스틱 / Plastics (PET)",
     "kw": ["플라스틱", "pet", "합성수지", "plastic", "수지"]},
    {"name": "알루미늄",     "code": "760110", "desc": "알루미늄 / Aluminium",
     "kw": ["알루미늄", "aluminum", "aluminium"]},
    # 의료/헬스
    {"name": "마스크",       "code": "630790", "desc": "마스크 / Face masks",
     "kw": ["마스크", "face mask", "surgical mask", "kf94", "kn95"]},
    {"name": "의약품",       "code": "300490", "desc": "의약품 / Medicaments",
     "kw": ["의약품", "약", "medicine", "pharmaceutical", "drug"]},
    {"name": "건강기능식품", "code": "210690", "desc": "건강기능식품 / Health supplements",
     "kw": ["건강기능식품", "영양제", "비타민", "health supplement", "vitamin", "프로바이오틱스", "probiotics"]},
    {"name": "의료기기",     "code": "901890", "desc": "의료기기 / Medical instruments",
     "kw": ["의료기기", "medical device", "혈압계", "혈당계", "체온계"]},
    # 가구/인테리어
    {"name": "가구",         "code": "940360", "desc": "가구 / Wooden furniture",
     "kw": ["가구", "의자", "책상", "테이블", "furniture", "chair", "desk", "table"]},
    {"name": "침구",         "code": "940490", "desc": "침구·매트리스 / Bedding & mattresses",
     "kw": ["침구", "이불", "베개", "매트리스", "bedding", "mattress", "pillow"]},
    {"name": "조명",         "code": "940540", "desc": "LED 조명 / LED lamps",
     "kw": ["조명", "led조명", "전구", "lighting", "led lamp", "bulb"]},
    # 스포츠/레저
    {"name": "골프용품",     "code": "950631", "desc": "골프채·골프용품 / Golf clubs & equipment",
     "kw": ["골프", "골프채", "골프공", "golf", "golf club", "golf ball"]},
    {"name": "자전거",       "code": "871200", "desc": "자전거 / Bicycles",
     "kw": ["자전거", "bicycle", "bike", "e-bike", "전기자전거"]},
    {"name": "운동기구",     "code": "950691", "desc": "헬스·피트니스 기구 / Fitness equipment",
     "kw": ["운동기구", "헬스기구", "트레드밀", "fitness equipment", "treadmill", "gym equipment"]},
    # 반려동물
    {"name": "반려동물용품", "code": "230910", "desc": "반려동물 사료·용품 / Pet food & products",
     "kw": ["반려동물", "펫", "사료", "강아지", "고양이", "pet food", "dog food", "cat food", "pet supplies"]},
    # 공예/생활
    {"name": "도자기",       "code": "691110", "desc": "도자기·자기 / Porcelain tableware",
     "kw": ["도자기", "자기", "도기", "ceramic", "porcelain", "pottery"]},
    {"name": "캔들",         "code": "340600", "desc": "양초·캔들 / Candles & tapers",
     "kw": ["캔들", "양초", "향초", "candle", "taper", "wax"]},
    {"name": "한복",         "code": "621140", "desc": "한복·전통의상 / Traditional garments",
     "kw": ["한복", "전통의상", "hanbok", "traditional clothing"]},
    {"name": "드론",         "code": "880211", "desc": "드론·무인항공기 / Unmanned aircraft (drones)",
     "kw": ["드론", "무인기", "무인항공기", "drone", "uav", "quadcopter"]},
    {"name": "주방용품",     "code": "732393", "desc": "냄비·팬·주방용품 / Cookware & kitchen items",
     "kw": ["냄비", "팬", "프라이팬", "주방용품", "cookware", "frying pan", "pot"]},
    {"name": "문구",         "code": "960200", "desc": "문구·사무용품 / Stationery",
     "kw": ["문구", "볼펜", "노트", "다이어리", "stationery", "pen", "notebook stationery"]},
    {"name": "완구",         "code": "950300", "desc": "장난감·완구 / Toys & games",
     "kw": ["장난감", "완구", "피규어", "toy", "figure", "doll", "게임"]},
    {"name": "악기",         "code": "920990", "desc": "악기·음악기기 / Musical instruments",
     "kw": ["악기", "기타", "피아노", "드럼", "musical instrument", "guitar", "piano"]},
    {"name": "수산물",       "code": "030290", "desc": "신선 수산물·어류 / Fresh fish & seafood",
     "kw": ["수산물", "생선", "어류", "참치", "연어", "fish", "seafood", "salmon", "tuna"]},
    {"name": "쌀",           "code": "100630", "desc": "쌀·백미 / Rice (milled)",
     "kw": ["쌀", "백미", "현미", "rice", "brown rice"]},
    {"name": "밀가루",       "code": "110100", "desc": "밀가루·소맥분 / Wheat flour",
     "kw": ["밀가루", "소맥분", "wheat flour", "flour"]},
    {"name": "화장솔",       "code": "960350", "desc": "화장용 솔·브러시 / Cosmetic brushes",
     "kw": ["화장솔", "메이크업브러시", "브러시", "cosmetic brush", "makeup brush"]},
    {"name": "네일",         "code": "330430", "desc": "네일 제품 / Nail preparations",
     "kw": ["네일", "매니큐어", "네일아트", "nail polish", "nail art", "manicure"]},
]

# UN Comtrade 국가 코드 → 한국어 이름
COUNTRY_NAMES = {
    "4": "아프가니스탄", "8": "알바니아", "12": "알제리", "24": "앙골라",
    "32": "아르헨티나", "36": "호주", "40": "오스트리아", "50": "방글라데시",
    "56": "벨기에", "64": "부탄", "76": "브라질", "116": "캄보디아",
    "124": "캐나다", "144": "스리랑카", "152": "칠레", "156": "중국",
    "170": "콜롬비아", "191": "크로아티아", "203": "체코", "208": "덴마크",
    "218": "에콰도르", "818": "이집트", "233": "에스토니아", "246": "핀란드",
    "251": "프랑스", "250": "프랑스", "276": "독일", "288": "가나", "300": "그리스",
    "344": "홍콩", "348": "헝가리", "356": "인도", "360": "인도네시아",
    "364": "이란", "368": "이라크", "372": "아일랜드", "376": "이스라엘",
    "380": "이탈리아", "388": "자메이카", "392": "일본", "400": "요르단",
    "404": "케냐", "410": "한국", "414": "쿠웨이트", "422": "레바논",
    "458": "말레이시아", "484": "멕시코", "504": "모로코", "104": "미얀마",
    "524": "네팔", "528": "네덜란드", "554": "뉴질랜드", "566": "나이지리아",
    "578": "노르웨이", "512": "오만", "586": "파키스탄", "604": "페루",
    "608": "필리핀", "616": "폴란드", "620": "포르투갈", "634": "카타르",
    "642": "루마니아", "643": "러시아", "682": "사우디아라비아", "710": "남아프리카공화국",
    "724": "스페인", "752": "스웨덴", "756": "스위스", "158": "대만",
    "764": "태국", "792": "터키", "784": "UAE", "826": "영국",
    "840": "미국", "858": "우루과이", "704": "베트남", "887": "예멘",
    # 추가 매핑
    "699": "인도", "702": "싱가포르", "703": "슬로바키아", "705": "슬로베니아",
    "96": "브루나이", "180": "콩고민주공화국", "192": "쿠바", "204": "베냉",
    "242": "피지", "418": "라오스", "508": "모잠비크", "132": "카보베르데",
    "20": "안도라",
}


# ──────────────────────────────────────────────
# 1단계: HS Code 후보 탐색 (키워드 매칭)
# ──────────────────────────────────────────────

def _sim(a: str, b: str) -> float:
    return SequenceMatcher(None, a, b).ratio()


def _find_hs_local(query: str, top_n: int = 3) -> list:
    query_l = query.lower().strip()
    results = []
    for item in HS_CODE_DB:
        best = 0
        if query_l == item["name"].lower():
            best = 100
        elif query_l in item["name"].lower() or item["name"].lower() in query_l:
            best = max(best, 80)
        for kw in item["kw"]:
            kw_l = kw.lower()
            if query_l == kw_l:
                best = max(best, 95)
            elif query_l in kw_l or kw_l in query_l:
                best = max(best, 75)
            else:
                best = max(best, int(_sim(query_l, kw_l) * 60))
        if best >= 25:
            results.append({**item, "match_score": best})
    results.sort(key=lambda x: x["match_score"], reverse=True)
    return results[:top_n]


def _find_hs_web(query: str) -> list:
    """로컬 DB 미스 시 DuckDuckGo로 HS코드 탐색"""
    import re
    try:
        from ddgs import DDGS
    except ImportError:
        return []

    # HS코드 패턴: 6자리 연속, 또는 xxxx.xx 형식
    hs_pat = re.compile(r'\b(\d{2})[.\s]?(\d{2})[.\s]?(\d{2})\b')
    # 10자리 한국 HSK 코드 (앞 6자리가 HS코드)
    hsk_pat = re.compile(r'\b(\d{4})[\. ](\d{2})\d{4}\b')

    queries = [
        f"{query} HS코드 품목번호 site:customs.go.kr OR site:tradenavi.or.kr",
        f"{query} HS code 수출 관세 품목분류",
        f"{query} HS code tariff number export classification",
    ]

    code_hits = {}
    code_desc = {}

    with DDGS() as ddgs:
        for q in queries:
            try:
                results = list(ddgs.text(q, max_results=6))
                time.sleep(0.4)
            except Exception:
                continue
            for r in results:
                text = r.get("title", "") + " " + r.get("body", "")
                title = r.get("title", "")

                # HS코드 관련 문맥 단어가 있는 경우만 유효 처리
                hs_context = bool(re.search(
                    r'HS|품목번호|관세|tariff|hscode|품목분류|hsn|hts', text, re.IGNORECASE
                ))
                if not hs_context:
                    continue

                # 일반 6자리 패턴
                for m in hs_pat.findall(text):
                    code = "".join(m)
                    ch = int(code[:2])
                    if ch == 0 or ch >= 98:
                        continue
                    code_hits[code] = code_hits.get(code, 0) + 1
                    if code not in code_desc:
                        code_desc[code] = title[:60]

                # HSK 10자리 → 앞 6자리 추출
                for m in hsk_pat.findall(text):
                    code = m[0] + m[1]
                    ch = int(code[:2])
                    if ch == 0 or ch >= 98:
                        continue
                    code_hits[code] = code_hits.get(code, 0) + 2
                    if code not in code_desc:
                        code_desc[code] = title[:60]

    if not code_hits:
        return []

    sorted_codes = sorted(code_hits.items(), key=lambda x: x[1], reverse=True)
    out = []
    for code, cnt in sorted_codes[:3]:
        out.append({
            "name": query,
            "code": code,
            "desc": f"{query} / {code_desc.get(code, code)}",
            "kw":   [query],
            "match_score": min(55 + cnt * 8, 78),
            "source": "web",
        })
    return out


def find_hs_codes(query: str, top_n: int = 3) -> list:
    local = _find_hs_local(query, top_n)

    # 로컬에서 충분히 확실한 결과가 있으면 바로 반환
    if local and local[0]["match_score"] >= 75:
        return local[:top_n]

    # 로컬 결과가 약하거나 없으면 웹 검색 병행
    web = _find_hs_web(query)
    # 이미 로컬에 있는 코드는 웹 결과에서 제외
    local_codes = {r["code"] for r in local}
    web_new = [r for r in web if r["code"] not in local_codes]

    merged = (local + web_new)
    merged.sort(key=lambda x: x["match_score"], reverse=True)
    return merged[:top_n]


# ──────────────────────────────────────────────
# 2단계: UN Comtrade Public API v1 호출
# 엔드포인트: https://comtradeapi.un.org/public/v1/preview
# 무료 플랜: 500 records/call, API 키 필요
# 키 발급: https://comtradedeveloper.un.org → 회원가입 → Subscribe
# ──────────────────────────────────────────────

COMTRADE_URL = "https://comtradeapi.un.org/public/v1/preview/C/A/HS"


def get_api_key() -> str:
    key = os.environ.get("COMTRADE_API_KEY", "").strip()
    if key:
        return key

    print("\n" + "─" * 60)
    print("  UN Comtrade API 키가 필요합니다.")
    print("  무료 발급 방법:")
    print("    1. https://comtradedeveloper.un.org 접속 → 회원가입")
    print("    2. APIs → Comtrade Preview → Subscribe")
    print("  또는 환경변수로 설정: export COMTRADE_API_KEY='your_key'")
    print("─" * 60)
    key = input("  API Key 입력 (건너뛰려면 Enter): ").strip()
    return key


def _fetch_one_year(hs_code: str, year: int, api_key: str) -> list:
    params = {
        "cmdCode": hs_code,
        "flowCode": "M",
        "partnerCode": "0",
        "period": str(year),
        "maxRecords": "500",
        "includeDesc": "true",
        "subscription-key": api_key,
    }
    try:
        resp = requests.get(COMTRADE_URL, params=params, timeout=30)
        if "text/html" in resp.headers.get("content-type", ""):
            print("    [오류] API 키 인증 실패")
            return []
        resp.raise_for_status()
        body = resp.json()
        return body.get("data", [])
    except requests.exceptions.HTTPError as e:
        print(f"    [오류] HTTP {e.response.status_code}")
    except requests.exceptions.Timeout:
        print("    [오류] 요청 타임아웃")
    except Exception as e:
        print(f"    [오류] {e}")
    return []


def _detect_latest_year(hs_code: str, api_key: str) -> int:
    """데이터가 존재하는 가장 최근 연도 탐지"""
    current_year = datetime.date.today().year
    for yr in range(current_year, current_year - 4, -1):
        params = {
            "cmdCode": hs_code, "flowCode": "M", "partnerCode": "0",
            "period": str(yr), "maxRecords": "1",
            "subscription-key": api_key,
        }
        try:
            r = requests.get(COMTRADE_URL, params=params, timeout=15)
            if r.headers.get("content-type", "").startswith("application/json"):
                if r.json().get("count", 0) > 0:
                    return yr
        except Exception:
            pass
    return current_year - 2  # 탐지 실패 시 2년 전으로 폴백


def fetch_trade_data(hs_code: str, api_key: str) -> dict:
    print("    최신 데이터 연도 확인 중...", end=" ", flush=True)
    latest = _detect_latest_year(hs_code, api_key)
    print(f"{latest}년")

    years = [latest, latest - 1]
    datasets = {}
    for i, yr in enumerate(years):
        print(f"    {yr}년 데이터 요청 중...", end=" ", flush=True)
        data = _fetch_one_year(hs_code, yr, api_key)
        print(f"{len(data)}건 수집")
        datasets[yr] = data
        if i < len(years) - 1:
            time.sleep(1)
    return datasets


# ──────────────────────────────────────────────
# 3단계: 스코어링
# ──────────────────────────────────────────────

def score_countries(datasets: dict) -> list:
    cur_yr = max(datasets.keys())
    prv_yr = min(datasets.keys())

    def to_dict(records):
        d = {}
        names = {}
        for r in records:
            code = str(r.get("reporterCode", "")).strip()
            val = r.get("primaryValue") or 0
            if code and val > 0:
                d[code] = d.get(code, 0) + val
                # reporterDesc가 있으면 한국어 이름 대신 사용 가능
                if code not in names:
                    names[code] = r.get("reporterDesc", "")
        return d, names

    cur, cur_names = to_dict(datasets[cur_yr])
    prv, _         = to_dict(datasets[prv_yr])

    if not cur:
        return []

    max_val = max(cur.values())
    rows = []

    for code, cur_val in cur.items():
        if code == "410":           # 한국 제외
            continue
        prv_val = prv.get(code, 0)
        growth = (cur_val - prv_val) / prv_val * 100 if prv_val > 0 else 0

        vol_score    = (cur_val / max_val) * 70
        growth_score = min(max((growth + 100) / 200 * 30, 0), 30)
        total        = vol_score + growth_score

        # 한국어 이름 우선, 없으면 API 응답 영어명 사용
        name = COUNTRY_NAMES.get(code) or cur_names.get(code, f"국가({code})")

        rows.append({
            "code": code,
            "name": name,
            "cur_val": cur_val,
            "prv_val": prv_val,
            "growth": growth,
            "vol_score": vol_score,
            "growth_score": growth_score,
            "total": total,
            "cur_yr": cur_yr,
            "prv_yr": prv_yr,
        })

    rows.sort(key=lambda x: x["total"], reverse=True)
    return rows


# ──────────────────────────────────────────────
# 경쟁국 분석 (수출 데이터, flowCode=X)
# ──────────────────────────────────────────────

KOREA_CODE = "410"


def fetch_export_data(hs_code: str, api_key: str) -> tuple:
    """전 세계 국가별 수출 데이터 조회 (경쟁국 파악용)"""
    current_year = datetime.date.today().year
    latest = current_year - 2
    print("    수출 데이터 최신 연도 확인 중...", end=" ", flush=True)
    for yr in range(current_year, current_year - 4, -1):
        params = {
            "cmdCode": hs_code, "flowCode": "X", "partnerCode": "0",
            "period": str(yr), "maxRecords": "1",
            "subscription-key": api_key,
        }
        try:
            r = requests.get(COMTRADE_URL, params=params, timeout=15)
            if r.headers.get("content-type", "").startswith("application/json"):
                if r.json().get("count", 0) > 0:
                    latest = yr
                    break
        except Exception:
            pass
    print(f"{latest}년")

    print(f"    {latest}년 수출 데이터 요청 중...", end=" ", flush=True)
    params = {
        "cmdCode": hs_code, "flowCode": "X", "partnerCode": "0",
        "period": str(latest), "maxRecords": "500",
        "includeDesc": "true", "subscription-key": api_key,
    }
    try:
        r = requests.get(COMTRADE_URL, params=params, timeout=30)
        data = r.json().get("data", [])
        print(f"{len(data)}건 수집")
        return data, latest
    except Exception as e:
        print(f"오류: {e}")
        return [], latest


def analyze_competitors(records: list, year: int) -> dict:
    """수출 데이터로 경쟁국 TOP 3 + 한국 현황 분석"""
    exports, names = {}, {}
    for r in records:
        code = str(r.get("reporterCode", ""))
        val  = r.get("primaryValue") or 0
        if code and val > 0:
            exports[code] = exports.get(code, 0) + val
            if code not in names:
                names[code] = r.get("reporterDesc", "")

    if not exports:
        return {}

    total      = sum(exports.values())
    korea_val  = exports.get(KOREA_CODE, 0)
    korea_share = korea_val / total * 100 if total > 0 else 0

    top3 = []
    for code, val in sorted(exports.items(), key=lambda x: x[1], reverse=True):
        if code == KOREA_CODE:
            continue
        share = val / total * 100
        if korea_val == 0:
            intensity = "측정불가"
        elif share >= korea_share * 3:
            intensity = "높음"
        elif share >= korea_share * 1.5:
            intensity = "중간"
        else:
            intensity = "낮음"
        top3.append({
            "code": code,
            "name": COUNTRY_NAMES.get(code) or names.get(code, f"국가({code})"),
            "export_val": val,
            "share": share,
            "intensity": intensity,
        })
        if len(top3) == 3:
            break

    all_sorted  = sorted(exports.items(), key=lambda x: x[1], reverse=True)
    korea_rank  = next((i + 1 for i, (c, _) in enumerate(all_sorted) if c == KOREA_CODE), None)

    return {
        "top3": top3,
        "korea": {"val": korea_val, "share": korea_share, "rank": korea_rank},
        "total": total,
        "year": year,
    }


def _extract_product_terms(hs_desc: str) -> list:
    """HS 설명에서 영문 검색 키워드 추출"""
    english = hs_desc.split("/")[-1].strip().lower()
    # 괄호 제거, 핵심 명사만
    import re
    english = re.sub(r"\(.*?\)", "", english).strip()
    terms = [english]
    # 2단어 이상이면 마지막 단어도 단독 추가 (예: "green tea" → "tea")
    words = english.split()
    if len(words) >= 2:
        terms.append(words[-1])
    return terms


def _is_relevant(result: dict, product_terms: list) -> bool:
    """결과 제목+스니펫에 제품 키워드가 하나라도 포함되는지 확인"""
    text = (result.get("title", "") + " " + result.get("body", "")).lower()
    return any(t in text for t in product_terms)


def _parse_results(results: list, product_terms: list, seen_domains: set,
                   limit: int, need_relevance: bool = True) -> list:
    from urllib.parse import urlparse
    skip = {"wikipedia.org", "youtube.com", "amazon.com", "alibaba.com",
            "zhihu.com", "reddit.com", "naver.com", "tistory.com", "blogspot.com"}
    found = []
    for r in results:
        if len(found) >= limit:
            break
        url   = r.get("href", "").strip()
        title = r.get("title", "").strip()
        body  = r.get("body", "").strip()
        if not title or not url:
            continue
        domain = urlparse(url).netloc.replace("www.", "")
        if domain in seen_domains:
            continue
        if any(s in domain for s in skip):
            continue
        if need_relevance and not _is_relevant(r, product_terms):
            continue
        seen_domains.add(domain)
        found.append({"name": title[:70], "url": url, "reason": body[:120].replace("\n", " ")})
    return found


def search_competitor_companies(country: str, hs_desc: str) -> list:
    """경쟁국 실제 수출 기업 검색 (제품 관련성 검증 포함)"""
    try:
        from ddgs import DDGS
    except ImportError:
        return []

    product_terms = _extract_product_terms(hs_desc)
    english_product = product_terms[0]

    queries = [
        f'"{english_product}" exporter {country} company brand manufacturer',
        f'site:kompass.com "{english_product}" {country}',
        f'site:europages.com "{english_product}" supplier {country}',
        f'{english_product} brand {country} export wholesale supplier',
    ]

    seen_domains = set()
    companies = []

    with DDGS() as ddgs:
        for q in queries:
            if len(companies) >= 2:
                break
            try:
                results = list(ddgs.text(q, max_results=6))
                time.sleep(0.8)
            except Exception:
                continue
            found = _parse_results(results, product_terms, seen_domains, 2 - len(companies))
            companies.extend(found)

    return companies


def print_competitor_analysis(result: dict):
    if not result or not result.get("top3"):
        return

    w  = 62
    kr = result["korea"]
    intensity_icon = {"높음": "⚠ 높음", "중간": "~ 중간", "낮음": "✓ 낮음"}

    print("\n" + "=" * w)
    print(f"  수출 경쟁사 분석  |  {result['year']}년 기준")
    print("=" * w)

    if kr["rank"] and kr["val"] > 0:
        print(f"\n  한국 현황  :  세계 {kr['rank']}위  |  "
              f"수출 {fmt_usd(kr['val'])}  |  점유율 {kr['share']:.1f}%")

    hs_desc = result.get("hs_desc", "")
    medals = ["1위", "2위", "3위"]
    for i, c in enumerate(result["top3"]):
        intensity = intensity_icon.get(c["intensity"], "")
        print(f"\n  {medals[i]}  {c['name']}  |  "
              f"수출 {fmt_usd(c['export_val'])}  |  점유율 {c['share']:.1f}%"
              + (f"  {intensity}" if intensity else ""))

        if hs_desc and c.get("companies"):
            for comp in c["companies"]:
                print(f"       - {comp['name']}")
                print(f"         {comp['url']}")

    print("=" * w)


# ──────────────────────────────────────────────
# 바이어 유형 추천 (룰 기반, 외부 API 없음)
# ──────────────────────────────────────────────

# 국가 그룹
_ADV = {"미국","영국","독일","프랑스","이탈리아","스페인","네덜란드","캐나다",
        "호주","일본","싱가포르","홍콩","스위스","스웨덴","덴마크","노르웨이",
        "오스트리아","벨기에","뉴질랜드","대만","이스라엘","핀란드","아일랜드"}
_SEA = {"베트남","태국","인도네시아","말레이시아","필리핀","캄보디아","미얀마","라오스","브루나이"}
_MID = {"UAE","사우디아라비아","카타르","쿠웨이트","오만","요르단","이집트","레바논"}
_SAS = {"인도","방글라데시","파키스탄","스리랑카","네팔"}

def _country_group(name: str) -> str:
    if name in _ADV: return "선진국"
    if name in _SEA: return "동남아"
    if name in _MID: return "중동"
    if name in _SAS: return "남아시아"
    return "기타신흥"

def _product_category(hs_code: str) -> str:
    return {
        "09":"식품음료","18":"식품음료","19":"식품음료","20":"식품음료",
        "21":"식품음료","22":"식품음료",
        "33":"화장품뷰티","34":"화장품뷰티",
        "30":"의약품",
        "61":"의류패션","62":"의류패션","63":"의류패션",
        "64":"신발잡화","42":"신발잡화",
        "84":"전자IT","85":"전자IT",
        "87":"자동차부품",
        "72":"산업소재","73":"산업소재","39":"산업소재",
    }.get(hs_code[:2], "일반상품")

# 채널 데이터: {카테고리: {국가그룹: [채널, ...]}}
_CH = {
    "식품음료": {
        "선진국": [
            {"type": "전문 아시아·건강식품 수입 유통사",
             "desc": "유기농·자연식품 전문 임포터 (유럽 Bio 인증 유통망 등)",
             "strategy": "유기농·ECOCERT 인증 취득 후 건강식품 전문 임포터와 독점 계약 추진"},
            {"type": "대형 슈퍼마켓·식료품 체인 바이어",
             "desc": "현지 대형 유통사 직매입 구매팀",
             "strategy": "KOTRA 해외무역관 바이어 매칭 활용, 판촉 행사 제안으로 입점 협상"},
            {"type": "온라인 이커머스 플랫폼",
             "desc": "Amazon, iHerb 등 글로벌 온라인 채널",
             "strategy": "Amazon FBA 또는 iHerb 브랜드 입점으로 소비자 직접 공략"},
        ],
        "동남아": [
            {"type": "마스터 디스트리뷰터",
             "desc": "현지 전국 유통망 보유 독점 총판",
             "strategy": "독점 계약 + 현지 마케팅 비용 분담으로 디스트리뷰터 파트너십 확보"},
            {"type": "대형 슈퍼마켓 체인",
             "desc": "Aeon, Big C, Lottemart 등 현지 대형마트",
             "strategy": "한류 연계 판촉·시식 행사 제안으로 대형마트 바이어 설득"},
            {"type": "크로스보더 이커머스",
             "desc": "Shopee, Lazada, Tokopedia 등 현지 플랫폼",
             "strategy": "K-푸드 프리미엄 이미지 + KOL 협업으로 온라인 선진입"},
        ],
        "중동": [
            {"type": "할랄 인증 식품 유통 에이전트",
             "desc": "할랄 인증 보유 식품 전문 유통상",
             "strategy": "할랄 인증 선취득 후 현지 에이전트 통해 B2B 납품 계약 추진"},
            {"type": "대형 하이퍼마켓 바이어",
             "desc": "Carrefour, LuLu Hypermarket 중동 구매팀",
             "strategy": "GCC 식품안전 인증(ESMA) 획득 후 중앙 구매팀에 직접 제안"},
            {"type": "온라인 식품 플랫폼",
             "desc": "Noon, Talabat Mart 등 현지 이커머스",
             "strategy": "아랍어 현지화 패키지로 온라인 채널 직접 입점"},
        ],
        "남아시아": [
            {"type": "지역별 식품 수입 에이전트",
             "desc": "주요 도시 기반 식품 수입 유통상",
             "strategy": "FSSAI 등 현지 식품 인허가 선취득 후 도시별 에이전트 계약 추진"},
            {"type": "모던 트레이드 슈퍼마켓 체인",
             "desc": "Reliance Fresh, Big Bazaar 등 현대식 유통",
             "strategy": "프리미엄·건강기능 포지셔닝으로 도시 중산층 타겟 바이어 접촉"},
            {"type": "이커머스 플랫폼",
             "desc": "Flipkart, Amazon India, JioMart 등",
             "strategy": "온라인 채널 선진입으로 브랜드 인지도 구축 후 오프라인 확장"},
        ],
        "기타신흥": [
            {"type": "식품 수입 에이전트",
             "desc": "현지 식품 수입 전문 에이전트",
             "strategy": "KOTRA 무역관 활용 에이전트 발굴, 독점 계약으로 진입 리스크 최소화"},
            {"type": "도매 식품 유통업체",
             "desc": "현지 B2B 도매 유통 네트워크",
             "strategy": "현지 식품 박람회 참가, 소량 샘플 오더로 관계 구축 후 물량 확대"},
            {"type": "SNS·이커머스 채널",
             "desc": "현지 SNS 커머스 및 이커머스 플랫폼",
             "strategy": "현지 인플루언서 마케팅 연계, 소비자 직접 판매 채널 우선 개설"},
        ],
    },
    "화장품뷰티": {
        "선진국": [
            {"type": "뷰티 전문 소매 체인 바이어",
             "desc": "Sephora, Douglas, Boots 등 현지 뷰티 체인",
             "strategy": "브랜드 피칭 + 팝업 스토어 운영으로 바이어 신뢰 확보 후 입점 제안"},
            {"type": "약국·파라파마시 유통사",
             "desc": "유럽 약국 체인 및 드럭스토어",
             "strategy": "더마 코스메틱 포지셔닝으로 약국 전문 유통사 계약 추진"},
            {"type": "뷰티 전문 이커머스",
             "desc": "Cult Beauty, Lookfantastic, Amazon Beauty",
             "strategy": "K-뷰티 트렌드 활용, 뷰티 플랫폼 입점 + 인플루언서 리뷰 마케팅"},
        ],
        "동남아": [
            {"type": "H&B 스토어·뷰티 멀티숍",
             "desc": "Watsons, Guardian, Boots Asia",
             "strategy": "K-뷰티 라인업으로 뷰티 체인 바이어 입점 제안, 한류 마케팅 연계"},
            {"type": "온라인 뷰티 플랫폼",
             "desc": "Shopee Beauty, Lazada Beauty, Sociolla",
             "strategy": "현지 뷰티 KOL 협업 + 플래시 세일로 빠른 인지도 확산"},
            {"type": "면세점·브랜드 직영 스토어",
             "desc": "공항 면세점 및 관광지 브랜드숍",
             "strategy": "면세점 입점으로 프리미엄 이미지 구축 후 일반 유통 확대"},
        ],
        "중동": [
            {"type": "뷰티·향수 전문 유통 에이전트",
             "desc": "GCC 뷰티 제품 전문 유통사",
             "strategy": "할랄 원료 인증 취득 + 에이전트 통해 럭셔리 뷰티 포지셔닝"},
            {"type": "대형 백화점·쇼핑몰 바이어",
             "desc": "Dubai Mall, Mall of Emirates 입점 바이어",
             "strategy": "Beautyworld Middle East 박람회 참가로 바이어 발굴"},
            {"type": "온라인 뷰티 플랫폼",
             "desc": "Noon Beauty, Namshi 등",
             "strategy": "아랍어 현지화 콘텐츠 + SNS 마케팅으로 온라인 채널 선점"},
        ],
        "남아시아": [
            {"type": "뷰티·퍼스널케어 유통사",
             "desc": "인도 뷰티 제품 전국 유통사",
             "strategy": "현지 유통사와 독점 계약, 도시 중산층 타겟 오프라인 판촉 전략"},
            {"type": "이커머스 플랫폼",
             "desc": "Nykaa, Amazon India, Flipkart",
             "strategy": "Nykaa 우선 입점, K-뷰티 카테고리 내 브랜드 인지도 구축"},
            {"type": "현대식 소매 체인",
             "desc": "Shoppers Stop, Lifestyle 등",
             "strategy": "프리미엄 포지셔닝으로 백화점 뷰티 섹션 입점, 체험 마케팅 진행"},
        ],
        "기타신흥": [
            {"type": "뷰티 수입 에이전트",
             "desc": "현지 화장품 수입 전문 에이전트",
             "strategy": "KOTRA 무역관 에이전트 발굴, K-뷰티 한류 마케팅 연계"},
            {"type": "H&B 스토어·드럭스토어",
             "desc": "현지 헬스앤뷰티 전문 소매 체인",
             "strategy": "샘플 제공 + 현지 뷰티 유튜버 협업으로 소비자 인지도 선확보"},
            {"type": "SNS·라이브 커머스",
             "desc": "Instagram, TikTok 기반 커머스",
             "strategy": "현지 뷰티 인플루언서 파트너십으로 SNS 커머스 채널 개설"},
        ],
    },
    "전자IT": {
        "선진국": [
            {"type": "전자제품 전문 유통사",
             "desc": "Media Markt, Best Buy, Currys 등",
             "strategy": "VAR 계약 추진, 기술 시연 및 A/S 보증으로 신뢰 확보"},
            {"type": "B2B IT 솔루션 바이어",
             "desc": "기업·공공기관 IT 구매 담당자",
             "strategy": "현지 시스템통합(SI) 업체 파트너십으로 B2B 대량 납품 채널 확보"},
            {"type": "온라인 이커머스",
             "desc": "Amazon, Best Buy Online 등",
             "strategy": "Amazon Vendor Central 직납 또는 FBA 활용, 리뷰 관리 집중"},
        ],
        "동남아": [
            {"type": "IT 제품 마스터 디스트리뷰터",
             "desc": "동남아 전자제품 전국 유통 총판",
             "strategy": "국가별 독점 디스트리뷰터 선정, 기술 교육 지원으로 판매 역량 강화"},
            {"type": "전자제품 전문 소매 체인",
             "desc": "Power Buy, iStudio, erafone 등",
             "strategy": "체험형 매장 내 시연 공간 확보 + 현지 프로모션 비용 지원"},
            {"type": "이커머스 플랫폼",
             "desc": "Shopee, Lazada, Tokopedia IT 카테고리",
             "strategy": "공식 브랜드 스토어 개설, 라이브 스트리밍으로 젊은 소비자층 공략"},
        ],
        "중동": [
            {"type": "IT 전문 유통 에이전트",
             "desc": "GCC IT 제품 전문 유통사",
             "strategy": "두바이 GITEX 박람회 참가로 중동 전역 유통 파트너 발굴"},
            {"type": "대형 하이퍼마켓·전자제품 체인",
             "desc": "LuLu Hypermarket, Sharaf DG 등",
             "strategy": "현지 대형 유통사 직납 계약 + 아랍어 제품 설명서 현지화 필수"},
            {"type": "정부·공공기관 조달 채널",
             "desc": "GCC 정부기관 IT 조달",
             "strategy": "현지 에이전트 통한 정부 입찰 참여, 현지 파트너십 의무화 대응"},
        ],
        "남아시아": [
            {"type": "IT 유통사 및 SI 업체",
             "desc": "Ingram Micro India 등 대형 IT 유통사",
             "strategy": "대형 IT 유통사 파트너십으로 전국 채널 확보, 기술 지원 차별화"},
            {"type": "전자제품 소매 체인",
             "desc": "Croma, Reliance Digital, Vijay Sales",
             "strategy": "대도시 프리미엄 소매 우선 입점, 브랜드 인지도 구축 후 확장"},
            {"type": "이커머스 플랫폼",
             "desc": "Flipkart, Amazon India",
             "strategy": "온라인 우선 진입, 플래시 세일로 빠른 시장 침투"},
        ],
        "기타신흥": [
            {"type": "IT 수입 에이전트",
             "desc": "현지 전자제품 수입 전문 에이전트",
             "strategy": "KOTRA 무역관 활용, 현지 인증 취득 선행 후 에이전트 계약"},
            {"type": "도매 전자제품 유통상",
             "desc": "현지 전자제품 도매 시장 유통상",
             "strategy": "현지 박람회 참가, 가격 경쟁력 및 A/S 정책 강조로 바이어 접촉"},
            {"type": "온라인 마켓플레이스",
             "desc": "현지 또는 글로벌 이커머스 플랫폼",
             "strategy": "온라인 채널로 먼저 브랜드 테스트 후 오프라인 파트너 발굴"},
        ],
    },
    "의류패션": {
        "선진국": [
            {"type": "패션 전문 에이전트·바이어",
             "desc": "국제 패션 박람회 기반 에이전트",
             "strategy": "Premiere Vision, Magic 등 국제 박람회 참가로 에이전트·바이어 직접 접촉"},
            {"type": "온라인 패션 플랫폼",
             "desc": "ASOS, Zalando, Farfetch 등",
             "strategy": "글로벌 패션 플랫폼 셀러 입점 + 트렌드 시즌 신상품 론칭 전략"},
            {"type": "대형 의류 유통·백화점 바이어",
             "desc": "백화점 및 대형 패션 유통 바이어",
             "strategy": "OEM/ODM 수주로 대형 브랜드 공급망 진입, 자체 브랜드 병행 전략"},
        ],
        "동남아": [
            {"type": "패션 마스터 디스트리뷰터",
             "desc": "동남아 패션 브랜드 총판",
             "strategy": "K-패션 트렌드 활용, 한류 콜라보로 현지 인지도 조기 확보"},
            {"type": "쇼핑몰·패션 소매 체인",
             "desc": "현지 쇼핑몰 패션 스트리트 입점사",
             "strategy": "팝업 스토어 운영으로 소비자 반응 테스트 후 정식 매장 결정"},
            {"type": "이커머스 패션 플랫폼",
             "desc": "Shopee Fashion, Lazada Fashion 등",
             "strategy": "한류 스타 착용 SNS 콘텐츠 연계, 이커머스 공식 스토어 운영"},
        ],
        "중동": [
            {"type": "패션·라이프스타일 유통 에이전트",
             "desc": "GCC 패션 유통 전문 에이전트",
             "strategy": "현지 복장 규정 고려한 디자인 현지화 + 모데스트 패션 라인 제안"},
            {"type": "럭셔리·프리미엄 쇼핑몰 입점",
             "desc": "Dubai Mall, Mall of the Emirates 입점",
             "strategy": "두바이 패션 위크 기간 팝업 부스 운영으로 프리미엄 이미지 구축"},
            {"type": "온라인 패션 플랫폼",
             "desc": "Namshi, Ounass, Noon Fashion",
             "strategy": "아랍어 현지화·사이즈 가이드 + 무료 반품 정책으로 온라인 구매 장벽 제거"},
        ],
        "남아시아": [
            {"type": "의류 도매 유통사",
             "desc": "인도 전국 의류 도매 네트워크",
             "strategy": "델리·뭄바이 의류 도매 시장 직접 방문, 현지 에이전트와 물량 협상"},
            {"type": "현대식 패션 소매 체인",
             "desc": "Pantaloons, Max Fashion, Westside 등",
             "strategy": "인도 중산층 타겟 가성비 라인업으로 패션 소매 체인 바이어 접촉"},
            {"type": "이커머스 패션 플랫폼",
             "desc": "Myntra, Nykaa Fashion, Flipkart Fashion",
             "strategy": "Myntra 우선 입점, 인도 트렌드 반영 컬렉션 기획"},
        ],
        "기타신흥": [
            {"type": "의류 수입 에이전트",
             "desc": "현지 의류 수입 전문 에이전트",
             "strategy": "현지 섬유·의류 박람회 참가, 에이전트 계약으로 초기 진입 비용 최소화"},
            {"type": "도매 의류 시장 유통상",
             "desc": "현지 의류 도매 시장 상인",
             "strategy": "가격 경쟁력 기본 아이템 중심으로 도매 바이어 접촉, 물량 확대 유도"},
            {"type": "SNS·이커머스 채널",
             "desc": "현지 SNS 커머스 및 이커머스",
             "strategy": "K-패션 콘텐츠로 SNS 팔로워 확보 후 이커머스 채널 연결"},
        ],
    },
    "자동차부품": {
        "선진국": [
            {"type": "OEM·Tier 1 부품 공급 바이어",
             "desc": "완성차 메이커 및 1차 협력사 구매팀",
             "strategy": "IATF 16949 인증 취득 후 OEM 공급망 진입, 품질 감사 철저 대비"},
            {"type": "애프터마켓 전문 유통사",
             "desc": "AutoZone, Bosch 유통망 등",
             "strategy": "가격·품질 경쟁력 갖춘 대체 부품으로 애프터마켓 유통사에 직납 제안"},
            {"type": "B2B 온라인 조달 플랫폼",
             "desc": "Würth, RS Components 등",
             "strategy": "B2B 플랫폼 입점으로 소량 다품종 대응, 빠른 납기 차별화"},
        ],
        "동남아": [
            {"type": "완성차 현지 조립공장 구매팀",
             "desc": "Toyota, Honda 등 현지 공장 구매팀",
             "strategy": "현지 완성차 공장 공급망 진입, 현지 법인 통한 안정적 납품 관계 구축"},
            {"type": "자동차 부품 총판·도매상",
             "desc": "현지 자동차 부품 전국 유통 총판",
             "strategy": "현지 총판 독점 계약, 기술 교육·재고 지원으로 파트너십 강화"},
            {"type": "자동차 정비소 네트워크",
             "desc": "현지 정비소 및 부품 소매점",
             "strategy": "정비소 네트워크 마케팅 + 기술 교육으로 실수요 창출"},
        ],
        "중동": [
            {"type": "자동차 딜러십 및 부품 에이전트",
             "desc": "GCC 공인 자동차 딜러 및 부품 에이전트",
             "strategy": "현지 공인 딜러십 파트너십으로 보증 수리 부품 납품 채널 확보"},
            {"type": "부품 도매 유통상",
             "desc": "두바이 자동차 부품 도매 시장",
             "strategy": "두바이 도매 시장 에이전트 활용, GCC 전역 유통망 확보"},
            {"type": "정부·공공기관 조달",
             "desc": "정부 차량 유지보수 조달 채널",
             "strategy": "현지 파트너사와 컨소시엄으로 정부 입찰 참여, 대형 계약 수주"},
        ],
        "남아시아": [
            {"type": "인도 완성차 공장 공급망",
             "desc": "Tata, Mahindra, Maruti 등 구매팀",
             "strategy": "Make in India 정책 활용, 현지 파트너와 합작으로 공급망 진입"},
            {"type": "자동차 부품 전국 유통사",
             "desc": "인도 부품 유통 대기업",
             "strategy": "인도 애프터마켓 규모 활용, 전국 유통망 보유 대형 딜러와 계약"},
            {"type": "이커머스 B2B 플랫폼",
             "desc": "IndiaMart, TradeIndia 등",
             "strategy": "IndiaMart 입점으로 소규모 정비소 수요 직접 공략, 빠른 납기 강조"},
        ],
        "기타신흥": [
            {"type": "자동차 부품 수입 에이전트",
             "desc": "현지 자동차 부품 수입 전문 에이전트",
             "strategy": "현지 에이전트 계약으로 수입 통관·인증 절차 위임, 초기 리스크 최소화"},
            {"type": "부품 도매 유통상",
             "desc": "현지 자동차 부품 도매 시장",
             "strategy": "현지 부품 박람회 참가, 가격 경쟁력 부각으로 도매 바이어 접촉"},
            {"type": "자동차 정비 프랜차이즈",
             "desc": "현지 자동차 정비 체인",
             "strategy": "정비 체인과 공급 계약으로 안정적 수요 확보, 정기 납품 시스템 구축"},
        ],
    },
}

_DEFAULT_CH = {
    "선진국": [
        {"type": "전문 수입 유통 에이전트",
         "desc": "현지 해당 품목 전문 수입 에이전트",
         "strategy": "KOTRA 해외무역관 통해 전문 에이전트 발굴 및 계약 추진"},
        {"type": "B2B 도매 유통상",
         "desc": "현지 도매 유통 네트워크",
         "strategy": "현지 무역 박람회 참가, 샘플 오더 우선 제안으로 관계 구축"},
        {"type": "온라인 B2B 플랫폼",
         "desc": "Alibaba, Global Sources 등",
         "strategy": "글로벌 B2B 플랫폼 프리미엄 셀러 등록으로 인바운드 바이어 유도"},
    ],
    "기타신흥": [
        {"type": "현지 수입 에이전트",
         "desc": "현지 수입 전문 에이전트",
         "strategy": "KOTRA 무역관 통해 현지 에이전트 발굴, 독점 계약으로 시장 진입"},
        {"type": "도매 유통상",
         "desc": "현지 도매 유통 네트워크",
         "strategy": "현지 박람회·무역 사절단 참가, 소량 공급으로 신뢰 구축 후 확대"},
        {"type": "온라인 마켓플레이스",
         "desc": "현지 이커머스 또는 글로벌 플랫폼",
         "strategy": "글로벌 B2B/B2C 플랫폼 활용으로 에이전트 없이 직접 시장 테스트"},
    ],
}


def get_buyer_channels(hs_code: str, target_country: str) -> list:
    cat   = _product_category(hs_code)
    group = _country_group(target_country)
    return (_CH.get(cat, {}).get(group)
            or _CH.get(cat, {}).get("기타신흥")
            or _DEFAULT_CH.get(group)
            or _DEFAULT_CH["기타신흥"])


def search_real_buyers(target_country: str, hs_desc: str) -> list:
    """DuckDuckGo 크롤링으로 실제 바이어 기업 검색 (제품 관련성 검증 포함)"""
    try:
        from ddgs import DDGS
    except ImportError:
        return []

    product_terms = _extract_product_terms(hs_desc)
    english_product = product_terms[0]

    queries = [
        f'"{english_product}" importer {target_country} wholesale company',
        f'site:kompass.com "{english_product}" importer {target_country}',
        f'site:europages.com "{english_product}" buyer {target_country}',
        f'{english_product} wholesale distributor {target_country} import company',
    ]

    seen_domains = set()
    buyers = []

    with DDGS() as ddgs:
        for q in queries:
            if len(buyers) >= 3:
                break
            try:
                results = list(ddgs.text(q, max_results=6))
                time.sleep(0.8)
            except Exception:
                continue
            found = _parse_results(results, product_terms, seen_domains, 3 - len(buyers))
            buyers.extend(found)

    return [{"name": b["name"], "url": b["url"], "reason": b.get("reason", "")} for b in buyers]


def print_buyer_analysis(channels: list, target_country: str, hs: dict, buyers: list = None):
    w = 62
    kotra_url = "https://www.kotra.or.kr/foreign/buyer/KTMITR060M.do"

    print("\n" + "=" * w)
    print(f"  바이어 추천  |  타겟 국가: {target_country}  |  품목: {hs['desc'].split('/')[0].strip()}")
    print("=" * w)

    # 실제 기업 검색 결과
    if buyers:
        print("\n  [ 실제 바이어 기업 ]")
        icons = ["①", "②", "③"]
        for i, b in enumerate(buyers[:3]):
            print(f"\n  {icons[i]} {b['name']}")
            print(f"     {b['url']}")
            if b["reason"]:
                print(f"     {b['reason'][:100]}")

    # 유통 채널 전략
    print(f"\n  [ 유통 채널 및 접근 전략 ]")
    ch_icons = ["▶", "▶", "▶"]
    for i, ch in enumerate(channels):
        print(f"\n  {ch_icons[i]} {ch['type']}")
        print(f"     → {ch['strategy']}")

    print(f"\n  {'─' * 58}")
    print(f"  KOTRA 바이어 DB  |  {target_country} 추가 바이어 검색")
    print(f"  {kotra_url}")
    print("=" * w)


# ──────────────────────────────────────────────
# 출력 유틸
# ──────────────────────────────────────────────

def fmt_usd(v: float) -> str:
    if v >= 1e9:
        return f"${v/1e9:.2f}B"
    if v >= 1e6:
        return f"${v/1e6:.1f}M"
    if v >= 1e3:
        return f"${v/1e3:.1f}K"
    return f"${v:.0f}"


def build_reason(c: dict, rank: int) -> str:
    labels = ["최우선 진출 추천 시장", "유망 진출 후보 시장", "성장 잠재력 보유 시장"]
    label = labels[rank - 1]
    name = c["name"]
    val_str = fmt_usd(c["cur_val"])
    g = c["growth"]

    line1 = f"[{label}] {name}의 {c['cur_yr']}년 수입 규모는 {val_str}입니다."

    if c["prv_val"] > 0:
        if g > 20:
            line2 = f"전년 대비 {g:.1f}% 급성장하며 수요가 빠르게 확대 중입니다."
        elif g > 5:
            line2 = f"전년 대비 {g:.1f}% 성장하며 안정적인 수요 증가세를 보입니다."
        elif g > -5:
            line2 = f"전년 대비 보합세({g:+.1f}%)로 안정적인 수입 시장을 유지합니다."
        else:
            line2 = f"전년 대비 {g:.1f}% 감소했으나 여전히 대규모 수입 시장입니다."
    else:
        line2 = f"전년도 비교 데이터가 없어 성장률은 산출되지 않았습니다."

    return f"{line1} {line2}"


def print_banner():
    w = 62
    print("\n" + "=" * w)
    print("        Global Trade Mentor AI  (MVP)")
    print("        수출 타겟 시장 분석 도구  |  UN Comtrade 기반")
    print("=" * w)


def print_candidates(candidates: list) -> dict:
    print("\n  [HS Code 후보 목록]")
    print("  " + "-" * 54)
    for i, c in enumerate(candidates, 1):
        print(f"  [{i}]  HS {c['code']}  {c['desc']}")
        print(f"        매칭 점수: {c['match_score']}점")
    print("  " + "-" * 54)

    while True:
        try:
            sel = input(f"\n  사용할 번호를 선택하세요 (1~{len(candidates)}): ").strip()
            idx = int(sel) - 1
            if 0 <= idx < len(candidates):
                return candidates[idx]
        except (ValueError, KeyboardInterrupt):
            pass
        print("  올바른 번호를 입력해주세요.")


def print_results(top3: list, hs: dict):
    w = 62
    medals = ["1위", "2위", "3위"]
    bars   = ["##########", "#######   ", "#####     "]

    print("\n" + "=" * w)
    print(f"  수출 타겟 국가 TOP 3  /  품목: {hs['desc']}")
    print(f"  HS Code: {hs['code']}  |  분석 기간: {top3[0]['prv_yr']}~{top3[0]['cur_yr']}")
    print("=" * w)

    for i, c in enumerate(top3):
        g_sign = "+" if c["growth"] >= 0 else ""
        g_str  = f"{g_sign}{c['growth']:.1f}%" if c["prv_val"] > 0 else "N/A"
        print(f"\n  {'★ ' if i==0 else '  '}{medals[i]}  {c['name']}")
        print(f"       수입 규모 : {fmt_usd(c['cur_val']):>10}   성장률: {g_str}")
        print(f"       종합 점수 : {c['total']:>6.1f}점   [{bars[i]}]")
        print()
        for line in _wrap(build_reason(c, i+1), 55):
            print(f"       {line}")

    print("\n" + "=" * w)
    print("  ※ UN Comtrade 공개 데이터 기반 / 실제 전략 수립 시")
    print("     추가 현장 조사 및 전문가 컨설팅을 병행하시기 바랍니다.")
    print("=" * w + "\n")


def _wrap(text: str, width: int) -> list:
    lines, cur = [], ""
    for word in text.split(" "):
        probe = (cur + " " + word).strip()
        if len(probe.encode("utf-8")) > width * 2:
            if cur:
                lines.append(cur)
            cur = word
        else:
            cur = probe
    if cur:
        lines.append(cur)
    return lines or [text]


# ──────────────────────────────────────────────
# Word 보고서 생성
# ──────────────────────────────────────────────

RANK_LABELS = ["최우선 진출 추천", "유망 진출 후보", "성장 잠재력 보유"]
RANK_COLORS = [RGBColor(0x1F, 0x49, 0x7D), RGBColor(0x2E, 0x74, 0xB5), RGBColor(0x5B, 0x9B, 0xD5)]


def _add_hyperlink(paragraph, url: str, display: str):
    """paragraph에 클릭 가능한 하이퍼링크 run 추가"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # 파란색 밑줄 스타일
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "맑은 고딕")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "21")  # 10.5pt = 21 half-points

    rPr.append(rFonts)
    rPr.append(color)
    rPr.append(u)
    rPr.append(sz)
    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = display
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def _set_font(run, size=11, bold=False, color=None):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    size = {1: 18, 2: 14, 3: 12}.get(level, 11)
    _set_font(run, size=size, bold=True, color=color or RGBColor(0x1F, 0x49, 0x7D))
    return p


def _body(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    _set_font(run, size=10.5)
    return p


def _add_country_table(doc, top3, cur_yr, prv_yr):
    """수입 데이터 요약 테이블"""
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["순위", "국가", f"수입규모\n({cur_yr})", f"성장률\n({prv_yr}→{cur_yr})", "종합점수"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        hdr_cells[i]._element.get_or_add_tcPr()
        # 헤더 배경색
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F497D")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)

    medals = ["🥇 1위", "🥈 2위", "🥉 3위"]
    for i, c in enumerate(top3):
        row = table.add_row().cells
        g_sign = "+" if c["growth"] >= 0 else ""
        g_str = f"{g_sign}{c['growth']:.1f}%" if c["prv_val"] > 0 else "N/A"
        vals = [medals[i], c["name"], fmt_usd(c["cur_val"]), g_str, f"{c['total']:.1f}점"]
        for j, v in enumerate(vals):
            p = row[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(v)
            _set_font(run, size=10.5, bold=(j == 1))


def generate_word_report(top3: list, hs: dict,
                         competitor: dict = None,
                         buyer_channels: list = None,
                         buyers: list = None,
                         output_path: str = None) -> str:
    doc = Document()

    # ── 페이지 여백 설정
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    today = datetime.date.today().strftime("%Y년 %m월 %d일")
    cur_yr = top3[0]["cur_yr"]
    prv_yr = top3[0]["prv_yr"]

    # ── 표지 영역
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("수출 타겟 시장 분석 보고서")
    _set_font(run, size=22, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Global Trade Mentor AI  |  MVP")
    _set_font(run, size=12, color=RGBColor(0x70, 0x70, 0x70))

    doc.add_paragraph()

    # ── 분석 개요 박스
    _heading(doc, "■ 분석 개요", level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    info_rows = [
        ("분석 품목",  hs["desc"]),
        ("HS Code",   hs["code"]),
        ("분석 기간", f"{prv_yr}년 ~ {cur_yr}년"),
        ("보고서 작성일", today),
    ]
    for row, (label, value) in zip(table.rows, info_rows):
        lp = row.cells[0].paragraphs[0]
        lr = lp.add_run(label)
        _set_font(lr, size=10.5, bold=True)
        vp = row.cells[1].paragraphs[0]
        vr = vp.add_run(value)
        _set_font(vr, size=10.5)

    doc.add_paragraph()

    # ── 분석 방법론
    _heading(doc, "■ 분석 방법론", level=2)
    _body(doc, (
        "본 보고서는 UN Comtrade 공개 무역 데이터를 기반으로 "
        f"HS Code {hs['code']} 품목의 국가별 수입 현황을 분석하였습니다. "
        "수입 규모(70%)와 전년 대비 성장률(30%)을 가중 합산하여 "
        "수출 유망 시장을 산정하였습니다."
    ))

    doc.add_paragraph()

    # ── 수입 데이터 요약 테이블
    _heading(doc, "■ 국가별 수입 현황 요약", level=2)
    _add_country_table(doc, top3, cur_yr, prv_yr)

    doc.add_paragraph()

    # ── TOP 3 상세 분석
    _heading(doc, "■ 수출 타겟 국가 TOP 3 상세 분석", level=2)

    for i, c in enumerate(top3):
        g_sign = "+" if c["growth"] >= 0 else ""
        g_str  = f"{g_sign}{c['growth']:.1f}%" if c["prv_val"] > 0 else "N/A"

        # 국가 제목
        p = doc.add_paragraph()
        run = p.add_run(f"{'★' if i == 0 else '◆'} {i+1}위  {c['name']}  —  {RANK_LABELS[i]} 시장")
        _set_font(run, size=12, bold=True, color=RANK_COLORS[i])

        # 수치 요약
        _body(doc, f"• 수입 규모 : {fmt_usd(c['cur_val'])}  ({cur_yr}년 기준)", indent=True)
        _body(doc, f"• 성장률   : {g_str}  ({prv_yr} → {cur_yr})", indent=True)
        _body(doc, f"• 종합 점수 : {c['total']:.1f}점  (수입규모 {c['vol_score']:.1f} + 성장률 {c['growth_score']:.1f})", indent=True)

        # 선정 근거
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run("[ 선정 근거 ]  ")
        _set_font(run, size=10.5, bold=True)
        run2 = p.add_run(build_reason(c, i + 1).split("] ", 1)[-1])
        _set_font(run2, size=10.5)

        if i < 2:
            doc.add_paragraph()

    doc.add_paragraph()

    # ── 경쟁국 분석
    if competitor and competitor.get("top3"):
        _heading(doc, "■ 수출 경쟁국 분석", level=2)
        kr = competitor["korea"]
        if kr["rank"] and kr["val"] > 0:
            _body(doc, f"한국 현황 : 세계 {kr['rank']}위  |  수출 {fmt_usd(kr['val'])}  |  점유율 {kr['share']:.1f}%  ({competitor['year']}년)")

        doc.add_paragraph()
        ctable = doc.add_table(rows=1, cols=4)
        ctable.style = "Table Grid"
        ctable.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, hdr in enumerate(["순위", "국가", "수출 규모", "경쟁 강도"]):
            p = ctable.rows[0].cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(hdr)
            _set_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tc_pr = ctable.rows[0].cells[idx]._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "1F497D"); shd.set(qn("w:val"), "clear")
            tc_pr.append(shd)
        intensity_map = {"높음": "⚠ 높음", "중간": "~ 중간", "낮음": "✓ 낮음"}
        for i, c in enumerate(competitor["top3"]):
            row = ctable.add_row().cells
            intensity_str = intensity_map.get(c["intensity"], "")
            for j, v in enumerate([f"{i+1}위", c["name"], fmt_usd(c["export_val"]), intensity_str]):
                p = row[j].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(v)
                _set_font(run, size=10.5, bold=(j == 1))

        # 경쟁사 기업명
        if any(c.get("companies") for c in competitor["top3"]):
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("주요 경쟁사")
            _set_font(run, size=11, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))
            for c in competitor["top3"]:
                if not c.get("companies"):
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(f"[ {c['name']} ]")
                _set_font(run, size=10.5, bold=True)
                for comp in c["companies"]:
                    p2 = doc.add_paragraph()
                    p2.paragraph_format.left_indent = Cm(1.0)
                    run2 = p2.add_run(f"• {comp['name']}  ")
                    _set_font(run2, size=10.5)
                    _add_hyperlink(p2, comp["url"], comp["url"])
        doc.add_paragraph()

    # ── 바이어 유형 추천
    if buyer_channels and top3:
        target_country = top3[0]["name"]
        _heading(doc, f"■ 바이어 추천  —  {target_country}", level=2)

        # 실제 바이어 기업
        if buyers:
            p = doc.add_paragraph()
            run = p.add_run("실제 바이어 기업")
            _set_font(run, size=11, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))
            for i, b in enumerate(buyers[:3]):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run_n = p.add_run(f"{'①②③'[i]}  {b['name']}  ")
                _set_font(run_n, size=10.5, bold=True)
                _add_hyperlink(p, b["url"], b["url"])
                if b.get("reason"):
                    _body(doc, b["reason"][:120], indent=True)
            doc.add_paragraph()

        # 유통 채널 전략
        p = doc.add_paragraph()
        run = p.add_run("유통 채널 및 접근 전략")
        _set_font(run, size=11, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))
        for i, ch in enumerate(buyer_channels):
            p = doc.add_paragraph()
            run = p.add_run(f"{'①②③'[i]}  {ch['type']}")
            _set_font(run, size=11, bold=True, color=RGBColor(0x2E, 0x74, 0xB5))
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(0.5)
            run2 = p2.add_run("→ 접근 전략: ")
            _set_font(run2, size=10.5, bold=True)
            run3 = p2.add_run(ch["strategy"])
            _set_font(run3, size=10.5)
            if i < 2:
                doc.add_paragraph()

        doc.add_paragraph()
        kotra_url = "https://www.kotra.or.kr/foreign/buyer/KTMITR060M.do"
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run_l = p.add_run(f"KOTRA 바이어 DB에서 {target_country} 추가 바이어를 검색하세요  →  ")
        _set_font(run_l, size=10.5, bold=True)
        _add_hyperlink(p, kotra_url, kotra_url)
        doc.add_paragraph()

    # ── 데이터 검수 링크
    _heading(doc, "■ 데이터 출처 및 직접 검수", level=2)

    # UN Comtrade 데이터 탐색기 링크 (해당 HS Code 직접 조회)
    explorer_url = (
        f"https://comtradeplus.un.org/TradeFlow"
        f"?Frequency=A&Flows=M&CommodityCodes={hs['code']}"
        f"&Partners=0&Reporters=all&period={cur_yr}&AggregateBy=none&TransportCode=total"
    )
    # UN Comtrade 전체 데이터베이스
    main_url = "https://comtradeplus.un.org"
    # API 문서
    api_doc_url = "https://comtradedeveloper.un.org"

    links = [
        (f"UN Comtrade 데이터 탐색기 — HS {hs['code']} {cur_yr}년 수입 데이터 직접 조회", explorer_url),
        ("UN Comtrade 데이터베이스 (전체 품목 검색)", main_url),
        ("UN Comtrade API 개발자 포털 (API 명세 및 쿼리 확인)", api_doc_url),
    ]

    for label, url in links:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        # 레이블 텍스트
        run_label = p.add_run(f"• {label}  →  ")
        _set_font(run_label, size=10.5)
        # 하이퍼링크 (XML 직접 삽입)
        _add_hyperlink(p, url, url)

    doc.add_paragraph()

    # ── 면책 조항
    p = doc.add_paragraph()
    run = p.add_run(
        "※ 본 보고서는 UN Comtrade 공개 데이터 기반의 자동 분석 결과이며, "
        "실제 수출 전략 수립 시 현지 시장 조사 및 전문가 컨설팅을 병행하시기 바랍니다."
    )
    _set_font(run, size=9, color=RGBColor(0x80, 0x80, 0x80))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 저장
    if not output_path:
        safe_name = hs["desc"].split("/")[0].strip().replace(" ", "_")
        output_path = f"report_{safe_name}_{today.replace(' ', '').replace('년','').replace('월','').replace('일','')}.docx"

    doc.save(output_path)
    return output_path


# ──────────────────────────────────────────────
# 메인
# ──────────────────────────────────────────────

def main():
    print_banner()

    # API 키 획득 (세션 전체에서 재사용)
    api_key = get_api_key()
    if not api_key:
        print("\n  API 키 없이는 실시간 데이터 수집이 불가합니다.")
        print("  https://comtradeplus.un.org 에서 무료 키를 발급받으세요.")
        sys.exit(1)

    print("\n  사용 가능 품목 예시: 스마트폰, 배터리, 라면, 화장품, 전기차, 철강 ...")

    while True:
        print()
        query = input("  물품명 입력 (종료: q) >>> ").strip()

        if query.lower() in ("q", "quit", "exit", "종료"):
            print("\n  프로그램을 종료합니다. 감사합니다!\n")
            sys.exit(0)

        if not query:
            continue

        # ── 1. HS Code 탐색 ──────────────────────
        print(f"\n  '{query}' 키워드 매칭 중...")
        candidates = find_hs_codes(query)

        if not candidates:
            print("  관련 HS Code를 찾지 못했습니다. 다른 키워드로 검색해보세요.")
            continue

        # ── 2. HS Code 선택 ──────────────────────
        selected = print_candidates(candidates)
        print(f"\n  선택된 HS Code: {selected['code']}  ({selected['desc']})")

        # ── 3. Comtrade+ API 데이터 수집 ─────────
        print("\n  UN Comtrade+ API 데이터 수집 중...")
        datasets = fetch_trade_data(selected["code"], api_key)

        cur_records = datasets.get(max(datasets.keys()), [])
        if not cur_records:
            print("\n  데이터를 가져오지 못했습니다.")
            print("  가능한 원인: API 키 오류 / 무료 플랜 일일 한도(250회) 초과 / 네트워크 오류")
            print("  잠시 후 다시 시도하거나 다른 HS Code를 선택해보세요.")
            continue

        # ── 4. 스코어링 ─────────────────────────
        print("  국가별 수입 규모 및 성장률 스코어링 중...")
        ranked = score_countries(datasets)

        if len(ranked) < 3:
            print(f"\n  충분한 데이터가 없습니다. (유효 국가: {len(ranked)}개)")
            continue

        # ── 5. 결과 출력 ─────────────────────────
        print_results(ranked[:3], selected)

        # ── 6. 경쟁사 분석 ───────────────────────
        print("\n  수출 경쟁사 분석 중...")
        export_records, export_year = fetch_export_data(selected["code"], api_key)
        competitor = analyze_competitors(export_records, export_year)
        competitor["hs_desc"] = selected["desc"]
        if competitor.get("top3"):
            for c in competitor["top3"]:
                print(f"    {c['name']} 경쟁사 검색 중...", end=" ", flush=True)
                c["companies"] = search_competitor_companies(c["name"], selected["desc"])
                print(f"{len(c['companies'])}개 발견")
        print_competitor_analysis(competitor)

        # ── 7. 바이어 추천 ───────────────────────
        target_country = ranked[0]["name"]
        channels = get_buyer_channels(selected["code"], target_country)
        print(f"\n  {target_country} 바이어 기업 검색 중...")
        buyers = search_real_buyers(target_country, selected["desc"])
        print_buyer_analysis(channels, target_country, selected, buyers)

        # ── 8. Word 보고서 자동 저장 ─────────────
        path = generate_word_report(ranked[:3], selected, competitor, channels, buyers)
        print(f"\n  보고서 저장 완료: {path}")
        import subprocess, platform
        if platform.system() == "Darwin":
            subprocess.Popen(["open", path])
        elif platform.system() == "Windows":
            subprocess.Popen(["start", path], shell=True)
        print()

        again = input("  다른 품목을 분석하시겠습니까? (y/n): ").strip().lower()
        if again != "y":
            print("\n  프로그램을 종료합니다. 감사합니다!\n")
            break


if __name__ == "__main__":
    main()
